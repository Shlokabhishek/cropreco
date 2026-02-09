"""
HTML to Word Document Converter
Converts HTML documents to Word (.docx) format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import re

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def set_document_styles(doc):
    """Set up custom styles for the document"""
    styles = doc.styles
    
    # Heading 1 style
    try:
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(26, 26, 26)
    except:
        pass
    
    # Heading 2 style
    try:
        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(44, 62, 80)
    except:
        pass
    
    # Heading 3 style
    try:
        h3_style = styles['Heading 3']
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(52, 73, 94)
    except:
        pass

def add_paragraph_with_style(doc, text, style=None, bold=False, italic=False, font_size=None):
    """Add a paragraph with specific styling"""
    if not text:
        return
    
    p = doc.add_paragraph(text, style=style)
    
    if bold or italic or font_size:
        run = p.runs[0] if p.runs else None
        if run:
            if bold:
                run.bold = bold
            if italic:
                run.italic = italic
            if font_size:
                run.font.size = Pt(font_size)
    
    return p

def add_table_from_html(doc, table_element):
    """Convert HTML table to Word table"""
    # Extract headers
    headers = [clean_text(th.get_text()) for th in table_element.find_all('th')]
    
    # Extract rows
    rows_data = []
    for tr in table_element.find_all('tr')[1:]:  # Skip header row
        row = [clean_text(td.get_text()) for td in tr.find_all('td')]
        if row:
            rows_data.append(row)
    
    if not headers or not rows_data:
        return
    
    # Create Word table
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Add data rows
    for row_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = str(cell_data)
                # Set font size for data
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    # Add spacing after table
    doc.add_paragraph()

def parse_html_to_docx(html_file, output_file):
    """Convert HTML file to Word document"""
    
    print(f"Converting {html_file} to {output_file}...")
    
    # Read HTML file
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    set_document_styles(doc)
    
    # Get and add title
    title_tag = soup.find('title')
    if title_tag:
        title_text = clean_text(title_tag.text)
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add meta info / author info
    meta_info = soup.find(class_=['meta-info', 'author-info', 'document-header'])
    if meta_info:
        meta_text = clean_text(meta_info.get_text())
        if meta_text:
            p = doc.add_paragraph(meta_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(11)
            doc.add_paragraph()  # Add spacing
    
    # Process abstract if present
    abstract = soup.find(class_=['abstract', 'executive-summary'])
    if abstract:
        # Add abstract heading
        abstract_heading = abstract.find(['h2', 'h3'])
        if abstract_heading:
            doc.add_heading(clean_text(abstract_heading.get_text()), level=2)
        
        # Add abstract paragraphs
        for p in abstract.find_all('p'):
            text = clean_text(p.get_text())
            if text:
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()  # Add spacing
    
    # Process main content
    processed_elements = set()
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table', 'div']):
        
        # Skip if already processed (e.g., inside abstract)
        if element in processed_elements or any(parent in processed_elements for parent in element.parents):
            continue
        
        # Skip style and script tags
        if element.name in ['style', 'script']:
            continue
        
        # Skip meta info and abstract (already processed)
        if 'meta-info' in element.get('class', []) or 'author-info' in element.get('class', []):
            continue
        if 'abstract' in element.get('class', []) or 'executive-summary' in element.get('class', []):
            continue
        
        # Process headings
        if element.name == 'h1':
            text = clean_text(element.get_text())
            if text and text != title_text:  # Skip duplicate title
                doc.add_heading(text, level=1)
                processed_elements.add(element)
        
        elif element.name == 'h2':
            text = clean_text(element.get_text())
            if text:
                doc.add_heading(text, level=2)
                processed_elements.add(element)
        
        elif element.name == 'h3':
            text = clean_text(element.get_text())
            if text:
                doc.add_heading(text, level=3)
                processed_elements.add(element)
        
        # Process paragraphs
        elif element.name == 'p':
            text = clean_text(element.get_text())
            if text and len(text) > 10:  # Skip very short texts
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Check for bold/italic
                if element.find('strong') or element.find('b'):
                    for run in para.runs:
                        run.bold = True
                if element.find('em') or element.find('i'):
                    for run in para.runs:
                        run.italic = True
                
                processed_elements.add(element)
        
        # Process lists
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = clean_text(li.get_text())
                if text:
                    doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
            processed_elements.add(element)
        
        # Process tables
        elif element.name == 'table':
            add_table_from_html(doc, element)
            processed_elements.add(element)
        
        # Process special divs (paper summaries, boxes, etc.)
        elif element.name == 'div':
            classes = element.get('class', [])
            
            # Paper summary boxes
            if 'paper-summary' in classes or 'paper-entry' in classes:
                doc.add_page_break()
                
                # Paper title
                paper_title = element.find(class_='paper-title')
                if paper_title:
                    text = clean_text(paper_title.get_text())
                    if text:
                        doc.add_heading(text, level=2)
                
                # Paper authors/meta
                paper_meta = element.find(class_=['paper-authors', 'paper-meta'])
                if paper_meta:
                    text = clean_text(paper_meta.get_text())
                    if text:
                        p = doc.add_paragraph(text)
                        for run in p.runs:
                            run.italic = True
                            run.font.size = Pt(10)
                
                # Process section boxes within paper
                for box in element.find_all(class_='section-box'):
                    box_title = box.find(class_='section-title')
                    if box_title:
                        text = clean_text(box_title.get_text())
                        if text:
                            doc.add_heading(text, level=3)
                    
                    # Add box content
                    for p in box.find_all('p'):
                        text = clean_text(p.get_text())
                        if text:
                            para = doc.add_paragraph(text)
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Add lists in box
                    for ul in box.find_all(['ul', 'ol']):
                        for li in ul.find_all('li', recursive=False):
                            text = clean_text(li.get_text())
                            if text:
                                doc.add_paragraph(text, style='List Bullet')
                
                processed_elements.add(element)
    
    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(output_file)
    print(f"✓ Successfully created {output_file}")
    
    # Count pages (approximate)
    total_paragraphs = len(doc.paragraphs)
    print(f"  Total paragraphs: {total_paragraphs}")

def main():
    """Main conversion function"""
    
    # List of files to convert
    files_to_convert = [
        ('RESEARCH_PAPER.html', 'RESEARCH_PAPER.docx'),
        ('LITERATURE_SURVEY.html', 'LITERATURE_SURVEY.docx'),
        ('RESEARCH_PAPERS_SUMMARY.html', 'RESEARCH_PAPERS_SUMMARY.docx'),
    ]
    
    print("=" * 60)
    print("HTML to Word Document Converter")
    print("=" * 60)
    
    for html_file, docx_file in files_to_convert:
        try:
            parse_html_to_docx(html_file, docx_file)
            print()
        except FileNotFoundError:
            print(f"✗ File not found: {html_file}")
            print()
        except Exception as e:
            print(f"✗ Error converting {html_file}: {str(e)}")
            import traceback
            traceback.print_exc()
            print()
    
    print("=" * 60)
    print("Conversion complete!")
    print("=" * 60)

if __name__ == "__main__":
    main()
